import os
import tempfile

from docx import Document

from docx_builder.figure import figure, figure_pair
from docx_builder.styles import StyleResolver

resolver = StyleResolver()
_FIG_STYLE = resolver.resolve("figure")
_FIG_PAIR_STYLE = resolver.resolve("figure_pair")
_CAPTION_STYLE = resolver.resolve("figure_caption")


def test_figure_missing_file_adds_placeholder() -> None:
    document = Document()
    figure(
        document,
        "missing.png",
        "Figure 1.1",
        "A caption",
        images_dir="/nonexistent_dir_xyz",
        figure_style=_FIG_STYLE,
        caption_style=_CAPTION_STYLE,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("[IMAGE NOT FOUND" in text for text in texts)


def test_figure_caption_label_appears() -> None:
    document = Document()
    figure(
        document,
        "missing.png",
        "Figure 2.3",
        "Caption text",
        images_dir="/nonexistent_dir_xyz",
        figure_style=_FIG_STYLE,
        caption_style=_CAPTION_STYLE,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("Figure 2.3" in text for text in texts)


def test_figure_with_real_image_adds_paragraphs() -> None:
    document = Document()
    with tempfile.TemporaryDirectory() as tmpdir:
        png_path = os.path.join(tmpdir, "sample.png")
        with open(png_path, "wb") as png_file:
            png_file.write(_minimal_png())
        initial_count = len(document.paragraphs)
        figure(
            document,
            "sample.png",
            "Figure 1.1",
            "Test caption",
            images_dir=tmpdir,
            figure_style=_FIG_STYLE,
            caption_style=_CAPTION_STYLE,
        )
        assert len(document.paragraphs) >= initial_count + 2


def test_figure_pair_missing_files_adds_caption() -> None:
    document = Document()
    figure_pair(
        document,
        "a.png",
        "b.png",
        "Figure 1.3",
        "Side-by-side",
        images_dir="/nonexistent_dir_xyz",
        figure_style=_FIG_PAIR_STYLE,
        caption_style=_CAPTION_STYLE,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("Figure 1.3" in text for text in texts)


def _minimal_png() -> bytes:
    return bytes(
        [
            0x89,
            0x50,
            0x4E,
            0x47,
            0x0D,
            0x0A,
            0x1A,
            0x0A,
            0x00,
            0x00,
            0x00,
            0x0D,
            0x49,
            0x48,
            0x44,
            0x52,
            0x00,
            0x00,
            0x00,
            0x01,
            0x00,
            0x00,
            0x00,
            0x01,
            0x08,
            0x02,
            0x00,
            0x00,
            0x00,
            0x90,
            0x77,
            0x53,
            0xDE,
            0x00,
            0x00,
            0x00,
            0x0C,
            0x49,
            0x44,
            0x41,
            0x54,
            0x08,
            0xD7,
            0x63,
            0xF8,
            0xCF,
            0xC0,
            0x00,
            0x00,
            0x00,
            0x02,
            0x00,
            0x01,
            0xE2,
            0x21,
            0xBC,
            0x33,
            0x00,
            0x00,
            0x00,
            0x00,
            0x49,
            0x45,
            0x4E,
            0x44,
            0xAE,
            0x42,
            0x60,
            0x82,
        ]
    )

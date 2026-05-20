from docx import Document
from docx.oxml.ns import qn

from docx_builder.table import remove_table_borders


def test_remove_table_borders_adds_borders_element() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    remove_table_borders(table)
    tbl_xml = table._tbl
    table_properties = tbl_xml.find(qn("w:tblPr"))
    assert table_properties is not None
    table_borders = table_properties.find(qn("w:tblBorders"))
    assert table_borders is not None


def test_remove_table_borders_all_sides_none() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    remove_table_borders(table)
    tbl_xml = table._tbl
    table_properties = tbl_xml.find(qn("w:tblPr"))
    assert table_properties is not None
    table_borders = table_properties.find(qn("w:tblBorders"))
    assert table_borders is not None
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_element = table_borders.find(qn(f"w:{side}"))
        assert border_element is not None
        assert border_element.get(qn("w:val")) == "none"

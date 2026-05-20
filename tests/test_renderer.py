import os
import tempfile

import pytest
from docx import Document

from docx_builder.renderer import fill_cover, render_sections
from docx_builder.styles import StyleResolver

_RESOLVER = StyleResolver()


def _minimal_png() -> bytes:
    return bytes(
        [
            0x89,
            0x50,
            0x4E,
            0x47,
            0x0D,
            0x0A,
            0x1A,
            0x0A,
            0x00,
            0x00,
            0x00,
            0x0D,
            0x49,
            0x48,
            0x44,
            0x52,
            0x00,
            0x00,
            0x00,
            0x01,
            0x00,
            0x00,
            0x00,
            0x01,
            0x08,
            0x02,
            0x00,
            0x00,
            0x00,
            0x90,
            0x77,
            0x53,
            0xDE,
            0x00,
            0x00,
            0x00,
            0x0C,
            0x49,
            0x44,
            0x41,
            0x54,
            0x08,
            0xD7,
            0x63,
            0xF8,
            0xCF,
            0xC0,
            0x00,
            0x00,
            0x00,
            0x02,
            0x00,
            0x01,
            0xE2,
            0x21,
            0xBC,
            0x33,
            0x00,
            0x00,
            0x00,
            0x00,
            0x49,
            0x45,
            0x4E,
            0x44,
            0xAE,
            0x42,
            0x60,
            0x82,
        ]
    )


_COVER_DATA: dict[str, object] = {
    "rows": ["Operating Systems", "CA1", "Mr. Smith", "Jane Doe", "2026001", "1 May 2026", "30 April 2026"],
    "ai_declaration": "I used AI.",
}


def test_fill_cover_fills_first_row() -> None:
    document = Document()
    document.add_table(rows=7, cols=2)
    fill_cover(document, _COVER_DATA)
    assert document.tables[0].rows[0].cells[1].paragraphs[0].text == "Operating Systems"


def test_fill_cover_fills_all_rows_by_index() -> None:
    document = Document()
    document.add_table(rows=3, cols=2)
    fill_cover(document, {"rows": ["Alpha", "Beta", "Gamma"]})
    table = document.tables[0]
    assert table.rows[0].cells[1].paragraphs[0].text == "Alpha"
    assert table.rows[1].cells[1].paragraphs[0].text == "Beta"
    assert table.rows[2].cells[1].paragraphs[0].text == "Gamma"


def test_fill_cover_variable_row_count() -> None:
    document = Document()
    document.add_table(rows=5, cols=2)
    fill_cover(document, {"rows": ["A", "B", "C", "D", "E"]})
    assert document.tables[0].rows[4].cells[1].paragraphs[0].text == "E"


def test_fill_cover_with_ai_declaration() -> None:
    document = Document()
    document.add_table(rows=1, cols=2)
    fill_cover(document, {"rows": ["X"], "ai_declaration": "I declare AI."})
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("I declare AI." in text for text in texts)


def test_fill_cover_without_ai_declaration() -> None:
    document = Document()
    document.add_table(rows=1, cols=2)
    fill_cover(document, {"rows": ["X"]})
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert not any("AI Use Declaration" in text for text in texts)


def test_fill_cover_ai_declaration_has_bold_prefix() -> None:
    document = Document()
    document.add_table(rows=7, cols=2)
    fill_cover(document, _COVER_DATA)
    declaration_paragraph = next(
        paragraph for paragraph in document.paragraphs if "AI Use Declaration" in paragraph.text
    )
    assert declaration_paragraph.runs[0].bold is True


def test_render_sections_h1() -> None:
    document = Document()
    render_sections(document, [{"call": "h1", "text": "My Heading"}], images_dir="/nonexistent", resolver=_RESOLVER)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "My Heading" in texts


def test_render_sections_h2() -> None:
    document = Document()
    render_sections(document, [{"call": "h2", "text": "Sub heading"}], images_dir="/nonexistent", resolver=_RESOLVER)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Sub heading" in texts


def test_render_sections_h3() -> None:
    document = Document()
    render_sections(document, [{"call": "h3", "text": "Sub sub"}], images_dir="/nonexistent", resolver=_RESOLVER)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Sub sub" in texts


def test_render_sections_body() -> None:
    document = Document()
    render_sections(
        document,
        [{"call": "body", "text": "Some body text."}],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Some body text." in texts


def test_render_sections_bullet() -> None:
    document = Document()
    render_sections(document, [{"call": "bullet", "text": "Item one"}], images_dir="/nonexistent", resolver=_RESOLVER)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("Item one" in text for text in texts)


def test_render_sections_bold_lead() -> None:
    document = Document()
    render_sections(
        document,
        [{"call": "bold_lead", "bold": "Key:", "rest": " explanation"}],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("Key:" in text for text in texts)


def test_render_sections_reference() -> None:
    document = Document()
    render_sections(
        document,
        [{"call": "reference", "text": "Smith (2020)"}],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Smith (2020)" in texts


def test_render_sections_page_break() -> None:
    document = Document()
    render_sections(document, [{"call": "page_break"}], images_dir="/nonexistent", resolver=_RESOLVER)
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "w:br" in all_xml


def test_render_sections_figure_missing_file() -> None:
    document = Document()
    render_sections(
        document,
        [{"call": "figure", "filename": "missing.png", "label": "Figure 1.1", "caption": "Cap"}],
        images_dir="/nonexistent_xyz",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("missing.png" in text for text in texts)


def test_render_sections_figure_existing_file() -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        png_path = os.path.join(tmpdir, "test.png")
        with open(png_path, "wb") as png_file:
            png_file.write(_minimal_png())
        document = Document()
        render_sections(
            document,
            [{"call": "figure", "filename": "test.png", "label": "Figure 1.1", "caption": "Cap"}],
            images_dir=tmpdir,
            resolver=_RESOLVER,
        )
        texts = [paragraph.text for paragraph in document.paragraphs]
        assert any("Figure 1.1" in text for text in texts)


def test_render_sections_figure_pair_missing_files() -> None:
    document = Document()
    render_sections(
        document,
        [
            {
                "call": "figure_pair",
                "filename1": "a.png",
                "filename2": "b.png",
                "label": "Figure 1.2",
                "caption": "Two images",
            }
        ],
        images_dir="/nonexistent_xyz",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert any("Figure 1.2" in text for text in texts)


def test_render_sections_multiple_calls() -> None:
    document = Document()
    render_sections(
        document,
        [
            {"call": "h1", "text": "Title"},
            {"call": "body", "text": "Content"},
            {"call": "reference", "text": "Ref (2020)"},
        ],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Title" in texts
    assert "Content" in texts
    assert "Ref (2020)" in texts


def test_render_sections_unknown_call_raises() -> None:
    document = Document()
    with pytest.raises(ValueError, match="unknown call"):
        render_sections(document, [{"call": "nonexistent"}], images_dir="/nonexistent", resolver=_RESOLVER)


def test_render_sections_toc_inserts_toc_field() -> None:
    document = Document()
    render_sections(document, [{"call": "toc", "levels": "1-2"}], images_dir="/nonexistent", resolver=_RESOLVER)
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "TOC" in all_xml


def test_render_sections_toc_default_levels() -> None:
    document = Document()
    render_sections(document, [{"call": "toc"}], images_dir="/nonexistent", resolver=_RESOLVER)
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "TOC" in all_xml


def test_render_sections_hide_page_counter_creates_section_break() -> None:
    document = Document()
    initial_sections = len(document.sections)
    render_sections(
        document,
        [
            {"call": "page_break", "hide_page_counter": True},
            {"call": "h1", "text": "Content"},
        ],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    assert len(document.sections) == initial_sections + 1


def test_render_sections_all_hidden_no_section_break() -> None:
    document = Document()
    initial_sections = len(document.sections)
    render_sections(
        document,
        [
            {"call": "h1", "text": "Title", "hide_page_counter": True},
            {"call": "body", "text": "Body", "hide_page_counter": True},
        ],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    assert len(document.sections) == initial_sections


def test_render_sections_no_hidden_no_section_break() -> None:
    document = Document()
    initial_sections = len(document.sections)
    render_sections(
        document,
        [{"call": "h1", "text": "Title"}, {"call": "body", "text": "Body"}],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    assert len(document.sections) == initial_sections


def test_render_sections_section_break_inserted_once() -> None:
    document = Document()
    render_sections(
        document,
        [
            {"call": "h1", "text": "Hidden", "hide_page_counter": True},
            {"call": "h1", "text": "Visible 1"},
            {"call": "h1", "text": "Visible 2"},
        ],
        images_dir="/nonexistent",
        resolver=_RESOLVER,
    )
    assert len(document.sections) == 2

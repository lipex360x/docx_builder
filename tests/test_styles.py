from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt, RGBColor

from docx_builder.builder import build
from docx_builder.styles import (
    StyleResolver,
    parse_align,
    parse_color,
    parse_length_inches,
    parse_length_pt,
)


def _write_yaml(path: Path, data: dict[str, Any]) -> None:
    with open(path, "w") as file:
        yaml.dump(data, file)


def test_parse_length_pt_int_string() -> None:
    assert parse_length_pt("11pt") == Pt(11)


def test_parse_length_pt_float_string() -> None:
    assert parse_length_pt("9.5pt") == Pt(9.5)


def test_parse_length_pt_bare_number_treated_as_pt() -> None:
    assert parse_length_pt("12") == Pt(12)


def test_parse_length_pt_none() -> None:
    assert parse_length_pt(None) is None


def test_parse_length_inches_in_suffix() -> None:
    assert parse_length_inches("0.3in") == parse_length_inches(0.3)


def test_parse_length_inches_cm_conversion() -> None:
    result = parse_length_inches("2.54cm")
    expected = parse_length_inches(1.0)
    assert result is not None and expected is not None
    assert abs(result.inches - expected.inches) < 0.001


def test_parse_color_with_hash() -> None:
    assert parse_color("#FF0033") == RGBColor(0xFF, 0x00, 0x33)


def test_parse_color_without_hash() -> None:
    assert parse_color("336699") == RGBColor(0x33, 0x66, 0x99)


def test_parse_align_valid() -> None:
    assert parse_align("center") is not None
    assert parse_align("justify") is not None


def test_resolver_returns_defaults_when_no_overrides() -> None:
    resolver = StyleResolver()
    style = resolver.resolve("h1")
    assert style["font_size"] == "14pt"
    assert style["bold"] is True


def test_resolver_global_override_replaces_default() -> None:
    resolver = StyleResolver(global_overrides={"h1": {"font_size": "20pt"}})
    style = resolver.resolve("h1")
    assert style["font_size"] == "20pt"
    assert style["bold"] is True


def test_resolver_inline_override_wins_over_global() -> None:
    resolver = StyleResolver(global_overrides={"h1": {"font_size": "20pt", "color": "#000000"}})
    style = resolver.resolve("h1", {"font_size": "24pt"})
    assert style["font_size"] == "24pt"
    assert style["color"] == "#000000"
    assert style["bold"] is True


def test_resolver_unknown_section_type_returns_overrides_only() -> None:
    resolver = StyleResolver(global_overrides={"custom": {"font_size": "10pt"}})
    style = resolver.resolve("custom")
    assert style == {"font_size": "10pt"}


def test_build_applies_global_h1_override(tmp_path: Path) -> None:
    content: dict[str, Any] = {
        "styles": {"h1": {"font_size": "20pt", "color": "#FF0000"}},
        "sections": [{"call": "h1", "text": "Big red title"}],
    }
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    document = Document(str(tmp_path / "Report.docx"))
    heading = next(p for p in document.paragraphs if p.text == "Big red title")
    run = heading.runs[0]
    assert run.font.size == Pt(20)
    assert run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_build_inline_override_per_section(tmp_path: Path) -> None:
    content: dict[str, Any] = {
        "styles": {"h1": {"font_size": "14pt"}},
        "sections": [
            {"call": "h1", "text": "Normal"},
            {"call": "h1", "text": "Special", "style": {"font_size": "22pt", "color": "#003366"}},
        ],
    }
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    document = Document(str(tmp_path / "Report.docx"))
    normal = next(p for p in document.paragraphs if p.text == "Normal")
    special = next(p for p in document.paragraphs if p.text == "Special")
    assert normal.runs[0].font.size == Pt(14)
    assert special.runs[0].font.size == Pt(22)
    assert special.runs[0].font.color.rgb == RGBColor(0x00, 0x33, 0x66)


def test_build_no_styles_block_keeps_defaults(tmp_path: Path) -> None:
    content = {"sections": [{"call": "h1", "text": "Default H1"}]}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    document = Document(str(tmp_path / "Report.docx"))
    heading = next(p for p in document.paragraphs if p.text == "Default H1")
    assert heading.runs[0].font.size == Pt(14)
    assert heading.runs[0].bold is True


def test_build_bullet_glyph_override(tmp_path: Path) -> None:
    content: dict[str, Any] = {
        "styles": {"bullet": {"glyph": "→ "}},
        "sections": [{"call": "bullet", "text": "Custom marker"}],
    }
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    document = Document(str(tmp_path / "Report.docx"))
    bullet_paragraph = next(p for p in document.paragraphs if "Custom marker" in p.text)
    assert "→ Custom marker" in bullet_paragraph.text

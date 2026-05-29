from __future__ import annotations

import subprocess
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from docx_builder import export
from docx_builder.builder import init_project
from docx_builder.summary import build_summary


def _make_docx_with_note(target: Path) -> Path:
    document = Document()
    document.add_paragraph("Heading")
    build_summary(document)
    document.add_paragraph(
        "Note: open in Microsoft Word, right-click the table of contents "
        "and select 'Update Field', or press Ctrl+A then F9."
    )
    document.add_paragraph("Body")
    document.save(str(target))
    return target


def test_resolve_input_uses_build_default_pattern(tmp_path: Path) -> None:
    init_project(tmp_path)

    resolved = export.resolve_input_path(tmp_path, None)

    assert resolved == tmp_path / "Report.docx"


def test_resolve_input_uses_cover_output_template(tmp_path: Path) -> None:
    (tmp_path / "content.yaml").write_text('cover:\n  output: "Doc_{number}.docx"\n  number: "0007"\n')

    resolved = export.resolve_input_path(tmp_path, None)

    assert resolved == tmp_path / "Doc_0007.docx"


def test_resolve_input_honours_explicit_override(tmp_path: Path) -> None:
    custom = tmp_path / "explicit.docx"

    resolved = export.resolve_input_path(tmp_path, str(custom))

    assert resolved == custom


def test_resolve_output_defaults_to_docx_with_pdf_suffix(tmp_path: Path) -> None:
    source = tmp_path / "Report_0001.docx"

    resolved = export.resolve_output_path(source, None)

    assert resolved == tmp_path / "Report_0001.pdf"


def test_resolve_output_honours_explicit_override(tmp_path: Path) -> None:
    source = tmp_path / "Report.docx"
    custom = tmp_path / "out" / "final.pdf"

    resolved = export.resolve_output_path(source, str(custom))

    assert resolved == custom


def test_strip_toc_note_removes_only_the_note_paragraph(tmp_path: Path) -> None:
    source = _make_docx_with_note(tmp_path / "with_note.docx")

    export.strip_toc_note(source)

    document = Document(str(source))
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert all(export.TOC_NOTE_NEEDLE not in text for text in texts)
    assert any("Heading" in text for text in texts)
    assert any("TABLE OF CONTENTS" in text for text in texts)


def test_export_rejects_non_macos(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "linux")
    source = _make_docx_with_note(tmp_path / "Report.docx")

    with pytest.raises(export.ExportError, match="requires macOS"):
        export.export_pdf(source, tmp_path / "Report.pdf")


def test_export_rejects_when_word_missing(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: False)
    source = _make_docx_with_note(tmp_path / "Report.docx")

    with pytest.raises(export.ExportError, match="Microsoft Word"):
        export.export_pdf(source, tmp_path / "Report.pdf")


def test_export_rejects_when_input_missing(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: True)
    missing = tmp_path / "absent.docx"

    with pytest.raises(export.ExportError, match=str(missing)):
        export.export_pdf(missing, tmp_path / "absent.pdf")


def _fake_word_run_factory(scratch_dir: Path, recorded: dict[str, Any], page_count: int = 3) -> Any:
    def fake_run(command: list[str], **_keyword_arguments: Any) -> subprocess.CompletedProcess[str]:
        name = Path(command[0]).name
        if name == "osascript":
            scratch_pdf = Path(command[-1])
            assert scratch_pdf.parent == scratch_dir
            scratch_docx = Path(command[-2])
            assert scratch_docx.parent == scratch_dir
            recorded["stripped_in_scratch"] = export.TOC_NOTE_NEEDLE not in _docx_text(scratch_docx)
            document = Document(str(scratch_docx))
            document.add_paragraph("WORD_POPULATED_TOC")
            document.save(str(scratch_docx))
            scratch_pdf.write_bytes(b"%PDF-1.4 fake")
            return subprocess.CompletedProcess(command, 0, "", "")
        if name == "mdimport":
            return subprocess.CompletedProcess(command, 0, "", "")
        if name == "mdls":
            return subprocess.CompletedProcess(command, 0, f"kMDItemNumberOfPages = {page_count}\n", "")
        raise AssertionError(f"unexpected command: {command}")

    return fake_run


def test_export_finalizes_source_by_default(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: True)

    scratch_dir = tmp_path / "scratch"
    monkeypatch.setattr(export, "SCRATCH_DIRECTORY", scratch_dir)

    source = _make_docx_with_note(tmp_path / "Report.docx")
    destination = tmp_path / "out" / "Report.pdf"

    recorded: dict[str, Any] = {}
    monkeypatch.setattr(subprocess, "run", _fake_word_run_factory(scratch_dir, recorded))

    export.export_pdf(source, destination)

    assert destination.exists()
    assert not list(scratch_dir.glob("*"))
    assert recorded["stripped_in_scratch"] is True
    source_text = _docx_text(source)
    assert export.TOC_NOTE_NEEDLE not in source_text
    assert "WORD_POPULATED_TOC" in source_text

    output = capsys.readouterr().out
    assert f"Exported: {destination} (3 pages)" in output


def test_export_no_update_source_leaves_source_untouched(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: True)

    scratch_dir = tmp_path / "scratch"
    monkeypatch.setattr(export, "SCRATCH_DIRECTORY", scratch_dir)

    source = _make_docx_with_note(tmp_path / "Report.docx")
    destination = tmp_path / "out" / "Report.pdf"
    original_bytes = source.read_bytes()

    recorded: dict[str, Any] = {}
    monkeypatch.setattr(subprocess, "run", _fake_word_run_factory(scratch_dir, recorded))

    export.export_pdf(source, destination, update_source=False)

    assert destination.exists()
    assert source.read_bytes() == original_bytes


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_parse_page_count_reads_mdls_value() -> None:
    assert export._parse_page_count("kMDItemNumberOfPages = 5\n") == 5


def test_parse_page_count_returns_none_when_absent() -> None:
    assert export._parse_page_count("kMDItemNumberOfPages = (null)\n") is None


def _fake_finalize_run_factory(scratch_dir: Path, recorded: dict[str, Any]) -> Any:
    def fake_run(command: list[str], **_keyword_arguments: Any) -> subprocess.CompletedProcess[str]:
        name = Path(command[0]).name
        if name == "osascript":
            scratch_docx = Path(command[-1])
            assert scratch_docx.parent == scratch_dir
            recorded["jxa_arguments"] = command[4:]
            recorded["stripped_in_scratch"] = export.TOC_NOTE_NEEDLE not in _docx_text(scratch_docx)
            document = Document(str(scratch_docx))
            document.add_paragraph("WORD_POPULATED_TOC")
            document.save(str(scratch_docx))
            return subprocess.CompletedProcess(command, 0, "", "")
        raise AssertionError(f"unexpected command: {command}")

    return fake_run


def test_finalize_source_writes_populated_docx_back(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: True)

    scratch_dir = tmp_path / "scratch"
    monkeypatch.setattr(export, "SCRATCH_DIRECTORY", scratch_dir)

    source = _make_docx_with_note(tmp_path / "Report.docx")

    recorded: dict[str, Any] = {}
    monkeypatch.setattr(subprocess, "run", _fake_finalize_run_factory(scratch_dir, recorded))

    export.finalize_source(source)

    assert not list(scratch_dir.glob("*"))
    assert recorded["stripped_in_scratch"] is True
    assert len(recorded["jxa_arguments"]) == 1
    source_text = _docx_text(source)
    assert export.TOC_NOTE_NEEDLE not in source_text
    assert "WORD_POPULATED_TOC" in source_text


def test_finalize_source_rejects_non_macos(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "linux")
    source = _make_docx_with_note(tmp_path / "Report.docx")

    with pytest.raises(export.ExportError, match="requires macOS"):
        export.finalize_source(source)


def test_finalize_source_rejects_when_word_missing(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setattr(export.sys, "platform", "darwin")
    monkeypatch.setattr(export, "_word_is_installed", lambda: False)
    source = _make_docx_with_note(tmp_path / "Report.docx")

    with pytest.raises(export.ExportError, match="Microsoft Word"):
        export.finalize_source(source)


@pytest.mark.requires_word
def test_export_pdf_real_word(tmp_path: Path) -> None:
    import os

    if os.environ.get("DOCX_BUILDER_TEST_WORD") != "1":
        pytest.skip("set DOCX_BUILDER_TEST_WORD=1 to run the real Word export test")

    source = _make_docx_with_note(tmp_path / "Report.docx")
    destination = tmp_path / "Report.pdf"

    export.export_pdf(source, destination)

    assert destination.exists()


@pytest.mark.requires_word
def test_finalize_source_real_word(tmp_path: Path) -> None:
    import os

    if os.environ.get("DOCX_BUILDER_TEST_WORD") != "1":
        pytest.skip("set DOCX_BUILDER_TEST_WORD=1 to run the real Word export test")

    source = _make_docx_with_note(tmp_path / "Report.docx")

    export.finalize_source(source)

    finalized = Document(str(source))
    assert not any("Note: open in Microsoft Word" in paragraph.text for paragraph in finalized.paragraphs)

from docx import Document

from docx_builder.summary import build_summary


def test_build_summary_toc_title() -> None:
    document = Document()
    build_summary(document)
    titles = [paragraph.text for paragraph in document.paragraphs]
    assert "TABLE OF CONTENTS" in titles


def test_build_summary_title_is_bold() -> None:
    document = Document()
    build_summary(document)
    toc_title = next(paragraph for paragraph in document.paragraphs if paragraph.text == "TABLE OF CONTENTS")
    assert toc_title.runs[0].bold is True


def test_build_summary_default_levels_in_xml() -> None:
    document = Document()
    build_summary(document)
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "1-2" in all_xml


def test_build_summary_custom_levels_in_xml() -> None:
    document = Document()
    build_summary(document, levels="1-3")
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "1-3" in all_xml


def test_build_summary_does_not_emit_f9_note() -> None:
    document = Document()
    build_summary(document)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert all("press Ctrl+A then F9" not in text for text in texts)
    assert all("Note: open in Microsoft Word" not in text for text in texts)


def test_build_summary_keeps_field_placeholder() -> None:
    document = Document()
    build_summary(document)
    all_xml = " ".join(paragraph._p.xml for paragraph in document.paragraphs)
    assert "Right-click here and select" in all_xml
    assert "TOC" in all_xml

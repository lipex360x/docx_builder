from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docx_builder.pagination import add_page_numbers
from docx_builder.styles import StyleResolver

_FOOTER_STYLE = StyleResolver().resolve("footer")


def test_single_section_has_page_field() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    footer_xml = document.sections[0].footer._element.xml
    assert "PAGE" in footer_xml


def test_single_section_has_section_pages_field() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    footer_xml = document.sections[0].footer._element.xml
    assert "SECTIONPAGES" in footer_xml


def test_footer_does_not_use_num_pages() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    footer_xml = document.sections[0].footer._element.xml
    assert "NUMPAGES" not in footer_xml


def test_single_section_footer_right_aligned() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    assert document.sections[0].footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_single_section_footer_font_size() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    footer = document.sections[0].footer
    assert any(run.font.size == Pt(9) for run in footer.paragraphs[0].runs)


def test_skip_cover_sections_cover_has_no_page_field() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    footer_xml = document.sections[0].footer._element.xml
    assert "PAGE" not in footer_xml


def test_skip_cover_sections_content_has_page_field() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    footer_xml = document.sections[1].footer._element.xml
    assert "PAGE" in footer_xml


def test_skip_cover_sections_content_has_section_pages_field() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    footer_xml = document.sections[1].footer._element.xml
    assert "SECTIONPAGES" in footer_xml


def test_skip_cover_sections_content_not_linked_to_previous() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    assert not document.sections[1].footer.is_linked_to_previous


def test_skip_cover_sections_cover_footer_is_empty() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    footer_xml = document.sections[0].footer._element.xml
    assert "NUMPAGES" not in footer_xml


def test_skip_cover_sections_restarts_body_page_numbering() -> None:
    document = Document()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    page_number_type = document.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert page_number_type is not None
    assert page_number_type.get(qn("w:start")) == "1"


def test_skip_cover_sections_overrides_existing_page_number_start() -> None:
    document = Document()
    document.add_section()
    existing = OxmlElement("w:pgNumType")
    existing.set(qn("w:start"), "9")
    document.sections[1]._sectPr.append(existing)
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    page_number_type = document.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert page_number_type is not None
    assert page_number_type.get(qn("w:start")) == "1"


def test_no_skip_does_not_restart_page_numbering() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE)
    assert document.sections[0]._sectPr.find(qn("w:pgNumType")) is None


def test_skip_cover_sections_single_section_does_not_restart() -> None:
    document = Document()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    assert document.sections[0]._sectPr.find(qn("w:pgNumType")) is None


def test_multiple_content_sections_all_have_footer() -> None:
    document = Document()
    document.add_section()
    document.add_section()
    add_page_numbers(document, style=_FOOTER_STYLE, skip_cover_sections=True)
    for section in document.sections[1:]:
        footer_xml = section.footer._element.xml
        assert "PAGE" in footer_xml

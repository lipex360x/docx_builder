from docx import Document
from docx.shared import Inches, Pt

from docx_builder.elements import body, bold_lead, bullet, heading_1, heading_2, heading_3, page_break, reference
from docx_builder.styles import StyleResolver

resolver = StyleResolver()


def test_page_break_adds_paragraph() -> None:
    document = Document()
    initial_count = len(document.paragraphs)
    page_break(document)
    assert len(document.paragraphs) == initial_count + 1


def test_h1_spacing() -> None:
    document = Document()
    result = heading_1(document, "Heading One", resolver.resolve("h1"))
    assert result.paragraph_format.space_before == Pt(14)
    assert result.paragraph_format.space_after == Pt(6)


def test_h1_run_font() -> None:
    document = Document()
    result = heading_1(document, "Heading One", resolver.resolve("h1"))
    assert result.runs[0].font.size == Pt(14)
    assert result.runs[0].bold is True


def test_h2_spacing_and_font() -> None:
    document = Document()
    result = heading_2(document, "Heading Two", resolver.resolve("h2"))
    assert result.paragraph_format.space_before == Pt(10)
    assert result.paragraph_format.space_after == Pt(3)
    assert result.runs[0].font.size == Pt(12)
    assert result.runs[0].bold is True


def test_h3_spacing_and_font() -> None:
    document = Document()
    result = heading_3(document, "Heading Three", resolver.resolve("h3"))
    assert result.paragraph_format.space_before == Pt(8)
    assert result.paragraph_format.space_after == Pt(3)
    assert result.runs[0].font.size == Pt(11)
    assert result.runs[0].bold is True


def test_body_spacing() -> None:
    document = Document()
    result = body(document, "Some paragraph text.", resolver.resolve("body"))
    assert result.paragraph_format.space_after == Pt(8)


def test_bullet_indent_and_marker() -> None:
    document = Document()
    result = bullet(document, "List item", resolver.resolve("bullet"))
    assert "• " in result.text
    assert result.paragraph_format.left_indent == Inches(0.3)
    assert result.paragraph_format.space_after == Pt(3)


def test_bold_lead_first_run_is_bold() -> None:
    document = Document()
    result = bold_lead(document, "Lead text: ", "regular text", resolver.resolve("bold_lead"))
    assert result.runs[0].bold is True
    assert not result.runs[1].bold


def test_bold_lead_indent() -> None:
    document = Document()
    result = bold_lead(document, "Lead: ", "rest", resolver.resolve("bold_lead"))
    assert result.paragraph_format.left_indent == Inches(0.3)
    assert result.paragraph_format.space_after == Pt(6)


def test_reference_hanging_indent() -> None:
    document = Document()
    result = reference(document, "Author (2024) Title. Publisher.", resolver.resolve("reference"))
    assert result.paragraph_format.left_indent == Inches(0.3)
    assert result.paragraph_format.first_line_indent == Inches(-0.3)
    assert result.paragraph_format.space_after == Pt(6)

from pathlib import Path

from docx import Document

from docx_builder.report import (
    DocumentReport,
    analyse,
    count_em_dashes,
    count_words,
    extract_text,
    format_report,
    reading_minutes,
    report_for,
)


def test_count_words_splits_on_whitespace() -> None:
    assert count_words("alpha beta gamma") == 3
    assert count_words("") == 0
    assert count_words("alpha   beta") == 2


def test_count_em_dashes_counts_only_em_dash() -> None:
    assert count_em_dashes("one — two — three") == 2
    assert count_em_dashes("one - two – three") == 0


def test_reading_minutes_divides_by_rate() -> None:
    assert reading_minutes(400) == 2.0
    assert reading_minutes(0) == 0.0
    assert reading_minutes(100, words_per_minute=50) == 2.0


def test_extract_text_includes_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("hello world")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Label")
    table.rows[0].cells[1].paragraphs[0].add_run("value — here")

    text = extract_text(document)

    assert "hello world" in text
    assert "value — here" in text


def test_analyse_counts_words_and_em_dashes() -> None:
    document = Document()
    document.add_paragraph("one two three — four")

    report = analyse(document)

    assert report.word_count == 5
    assert report.em_dash_count == 1


def test_report_for_reads_saved_docx(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("alpha beta — gamma")
    saved = tmp_path / "Report.docx"
    document.save(str(saved))

    report = report_for(saved)

    assert report.word_count == 4
    assert report.em_dash_count == 1


def test_format_report_contains_all_metrics() -> None:
    report = DocumentReport(word_count=120, em_dash_count=0, reading_minutes=2.0)

    rendered = format_report(report)

    assert "Words: 120" in rendered
    assert "Reading time:" in rendered
    assert "Pages:" in rendered
    assert "Em-dashes (U+2014): 0" in rendered


def test_format_report_flags_em_dashes_when_present() -> None:
    report = DocumentReport(word_count=120, em_dash_count=3, reading_minutes=2.0)

    rendered = format_report(report)

    assert "Em-dashes (U+2014): 3" in rendered
    assert "remove" in rendered.lower()


def test_format_report_does_not_flag_when_zero() -> None:
    report = DocumentReport(word_count=120, em_dash_count=0, reading_minutes=2.0)

    rendered = format_report(report)

    assert "remove" not in rendered.lower()


def test_format_report_reading_under_one_minute() -> None:
    report = DocumentReport(word_count=10, em_dash_count=0, reading_minutes=0.05)

    rendered = format_report(report)

    assert "<1 min" in rendered


def test_format_report_reading_rounds_minutes() -> None:
    report = DocumentReport(word_count=600, em_dash_count=0, reading_minutes=3.0)

    rendered = format_report(report)

    assert "~3 min" in rendered

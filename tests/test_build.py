import os
import shutil
from pathlib import Path

import yaml
from docx import Document

from docx_builder.builder import build, has_toc

_MINIMAL_COVER: dict[str, object] = {
    "template": "Cover.docx",
    "number": "9999999",
    "rows": ["Test Module", "CA1", "Mr. Test", "Test Student", "9999999", "1 May 2026", "30 April 2026"],
    "ai_declaration": "I used AI.",
}


def _write_yaml(path: Path, data: dict[str, object]) -> None:
    with open(path, "w") as content_file:
        yaml.dump(data, content_file)


def test_has_toc_detects_toc_in_front_matter(tmp_path: Path) -> None:
    content = {"front_matter": [{"call": "toc", "levels": "1-2"}], "sections": [{"call": "h1", "text": "X"}]}
    _write_yaml(tmp_path / "content.yaml", content)

    assert has_toc(tmp_path) is True


def test_has_toc_detects_toc_in_sections(tmp_path: Path) -> None:
    content = {"sections": [{"call": "toc"}, {"call": "h1", "text": "X"}]}
    _write_yaml(tmp_path / "content.yaml", content)

    assert has_toc(tmp_path) is True


def test_has_toc_false_when_no_toc(tmp_path: Path) -> None:
    content = {"sections": [{"call": "h1", "text": "X"}, {"call": "body", "text": "Y"}]}
    _write_yaml(tmp_path / "content.yaml", content)

    assert has_toc(tmp_path) is False


def test_build_creates_docx(tmp_path: Path, cover_dir: Path) -> None:
    content = {"cover": _MINIMAL_COVER, "sections": [{"call": "h1", "text": "Introduction"}]}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_without_cover_creates_blank_document(tmp_path: Path) -> None:
    content = {"sections": [{"call": "h1", "text": "Title only"}]}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    assert (tmp_path / "Report.docx").exists()


def test_build_cover_without_template_uses_blank_document(tmp_path: Path) -> None:
    cover = {key: value for key, value in _MINIMAL_COVER.items() if key != "template"}
    content = {"cover": cover, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_output_template_pattern(tmp_path: Path, cover_dir: Path) -> None:
    cover = {**_MINIMAL_COVER, "output": "MyReport_{number}.docx"}
    content = {"cover": cover, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    assert (tmp_path / "MyReport_9999999.docx").exists()


def test_build_output_override_wins(tmp_path: Path, cover_dir: Path) -> None:
    content = {"cover": _MINIMAL_COVER, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir), output_override="custom.docx")

    assert (tmp_path / "custom.docx").exists()


def test_build_template_relative_to_project_dir(tmp_path: Path, cover_docx: Path) -> None:
    templates_dir = tmp_path / ".." / "shared_templates"
    templates_dir.mkdir()
    shutil.copy(str(cover_docx), templates_dir / "Cover.docx")
    cover = {**_MINIMAL_COVER, "template": "../shared_templates/Cover.docx"}
    content = {"cover": cover, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path))

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_uses_template_from_yaml(tmp_path: Path, cover_docx: Path) -> None:
    alt_cover = os.path.join(str(tmp_path), "AltCover.docx")
    shutil.copy(str(cover_docx), alt_cover)
    template_dir = str(tmp_path)
    cover = {**_MINIMAL_COVER, "template": "AltCover.docx"}
    content = {"cover": cover, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=template_dir)

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_without_ai_declaration(tmp_path: Path, cover_dir: Path) -> None:
    cover = {key: value for key, value in _MINIMAL_COVER.items() if key != "ai_declaration"}
    content = {"cover": cover, "sections": []}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_with_all_section_types(tmp_path: Path, cover_dir: Path) -> None:
    content = {
        "cover": _MINIMAL_COVER,
        "sections": [
            {"call": "page_break", "hide_page_counter": True},
            {"call": "toc", "levels": "1-2", "hide_page_counter": True},
            {"call": "h1", "text": "Part 1"},
            {"call": "h2", "text": "Section 1.1"},
            {"call": "body", "text": "Body text."},
            {"call": "bullet", "text": "Item"},
            {"call": "bold_lead", "bold": "Key:", "rest": " value"},
            {"call": "reference", "text": "Smith (2020)"},
            {"call": "page_break"},
            {"call": "figure", "filename": "missing.png", "label": "Figure 1.1", "caption": "Cap"},
        ],
    }
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    assert (tmp_path / "Report_9999999.docx").exists()


def test_build_no_hidden_sections_has_page_numbers_everywhere(tmp_path: Path, cover_dir: Path) -> None:
    content = {"cover": _MINIMAL_COVER, "sections": [{"call": "h1", "text": "Introduction"}]}
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    output_document = Document(str(tmp_path / "Report_9999999.docx"))
    footer_xml = output_document.sections[0].footer._element.xml
    assert "PAGE" in footer_xml


def test_build_front_matter_suppresses_footer(tmp_path: Path) -> None:
    content = {
        "front_matter": [
            {"call": "page_break"},
            {"call": "toc", "levels": "1-2"},
        ],
        "sections": [
            {"call": "h1", "text": "Introduction"},
        ],
    }
    _write_yaml(tmp_path / "content.yaml", content)

    build(str(tmp_path))

    output_document = Document(str(tmp_path / "Report.docx"))
    assert len(output_document.sections) >= 2
    front_footer_xml = output_document.sections[0].footer._element.xml
    body_footer_xml = output_document.sections[1].footer._element.xml
    assert "PAGE" not in front_footer_xml
    assert "PAGE" in body_footer_xml


def test_build_page_numbers_disabled_clears_all_footers(tmp_path: Path) -> None:
    content = {
        "page_numbers": False,
        "sections": [
            {"call": "h1", "text": "Curriculum"},
            {"call": "body", "text": "One-page resume content."},
        ],
    }
    _write_yaml(tmp_path / "content.yaml", content)

    build(str(tmp_path))

    output_document = Document(str(tmp_path / "Report.docx"))
    for section in output_document.sections:
        footer_xml = section.footer._element.xml
        assert "PAGE" not in footer_xml


def test_build_hide_page_counter_on_sections_suppresses_cover_footer(tmp_path: Path, cover_dir: Path) -> None:
    content = {
        "cover": _MINIMAL_COVER,
        "sections": [
            {"call": "page_break", "hide_page_counter": True},
            {"call": "toc", "levels": "1-2", "hide_page_counter": True},
            {"call": "h1", "text": "Introduction"},
        ],
    }
    _write_yaml(tmp_path / "content.yaml", content)
    (tmp_path / "images").mkdir()

    build(str(tmp_path), template_dir=str(cover_dir))

    output_document = Document(str(tmp_path / "Report_9999999.docx"))
    assert len(output_document.sections) >= 2
    cover_footer_xml = output_document.sections[0].footer._element.xml
    assert "PAGE" not in cover_footer_xml
    content_footer_xml = output_document.sections[1].footer._element.xml
    assert "PAGE" in content_footer_xml

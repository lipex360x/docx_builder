from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType

WORDS_PER_MINUTE = 200
EM_DASH = "—"


@dataclass(frozen=True)
class DocumentReport:
    word_count: int
    em_dash_count: int
    reading_minutes: float


def extract_text(document: DocumentType) -> str:
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def count_words(text: str) -> int:
    return len(text.split())


def count_em_dashes(text: str) -> int:
    return text.count(EM_DASH)


def reading_minutes(word_count: int, words_per_minute: int = WORDS_PER_MINUTE) -> float:
    return word_count / words_per_minute


def analyse(document: DocumentType) -> DocumentReport:
    text = extract_text(document)
    word_count = count_words(text)
    return DocumentReport(
        word_count=word_count,
        em_dash_count=count_em_dashes(text),
        reading_minutes=reading_minutes(word_count),
    )


def report_for(docx_path: str | Path) -> DocumentReport:
    return analyse(Document(str(docx_path)))


def format_report(report: DocumentReport) -> str:
    reading = "<1 min" if report.reading_minutes < 1 else f"~{round(report.reading_minutes)} min"
    flag = "  <- remove before shipping" if report.em_dash_count else ""
    lines = [
        "Report:",
        f"  Words: {report.word_count}",
        f"  Reading time: {reading}",
        "  Pages: n/a (known only after export to PDF)",
        f"  Em-dashes (U+2014): {report.em_dash_count}{flag}",
    ]
    return "\n".join(lines)

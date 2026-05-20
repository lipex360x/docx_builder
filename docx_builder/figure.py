from __future__ import annotations

import os
from typing import Any

from docx.document import Document
from docx.shared import Inches

from docx_builder.styles import apply_paragraph_style, apply_run_style, parse_align, parse_length_inches
from docx_builder.table import remove_table_borders


def _resolve_width(value: object, fallback_inches: float) -> Inches:
    if value is None:
        return Inches(fallback_inches)
    parsed = parse_length_inches(value)
    if parsed is None:
        return Inches(fallback_inches)
    return parsed


def figure(
    document: Document,
    filename: str,
    label: str,
    caption: str,
    *,
    images_dir: str,
    figure_style: dict[str, Any],
    caption_style: dict[str, Any],
    width: float | str | None = None,
) -> None:
    image_path = os.path.join(images_dir, filename)
    if not os.path.exists(image_path):
        document.add_paragraph(f"[IMAGE NOT FOUND: {filename}]")
        _add_caption(document, label, caption, caption_style)
        return
    image_paragraph = document.add_paragraph()
    apply_paragraph_style(image_paragraph, figure_style)
    if width is None:
        width = figure_style.get("default_width", 5.0)
    image_paragraph.add_run().add_picture(image_path, width=_resolve_width(width, 5.0))
    _add_caption(document, label, caption, caption_style)


def figure_pair(
    document: Document,
    filename1: str,
    filename2: str,
    label: str,
    caption: str,
    *,
    images_dir: str,
    figure_style: dict[str, Any],
    caption_style: dict[str, Any],
    width1: float | str | None = None,
    width2: float | str | None = None,
) -> None:
    image_table = document.add_table(rows=1, cols=2)
    remove_table_borders(image_table)
    image_table.autofit = False
    if width1 is None:
        width1 = figure_style.get("default_width_left", 3.0)
    if width2 is None:
        width2 = figure_style.get("default_width_right", 2.5)
    pairs: list[tuple[str, float | str]] = [(filename1, width1), (filename2, width2)]
    for index, (current_filename, current_width) in enumerate(pairs):
        image_path = os.path.join(images_dir, current_filename)
        cell = image_table.cell(0, index)
        cell_paragraph = cell.paragraphs[0]
        apply_paragraph_style(cell_paragraph, figure_style)
        align = parse_align(figure_style.get("align"))
        if align is not None:
            cell_paragraph.alignment = align
        if os.path.exists(image_path):
            cell_paragraph.add_run().add_picture(image_path, width=_resolve_width(current_width, 3.0))
        else:
            cell_paragraph.add_run(f"[IMAGE NOT FOUND: {current_filename}]")
    _add_caption(document, label, caption, caption_style)


def _add_caption(document: Document, label: str, caption: str, style: dict[str, Any]) -> None:
    caption_paragraph = document.add_paragraph()
    apply_paragraph_style(caption_paragraph, style)

    label_text = f"{label} – "
    run_label = caption_paragraph.add_run(label_text)
    apply_run_style(run_label, style)
    if style.get("label_bold", True):
        run_label.bold = True

    run_caption = caption_paragraph.add_run(caption)
    apply_run_style(run_caption, style)
    if style.get("caption_italic", True):
        run_caption.italic = True

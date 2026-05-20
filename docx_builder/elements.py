from __future__ import annotations

from typing import Any, cast

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from docx_builder.styles import apply_paragraph_style, apply_run_style


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    break_element = OxmlElement("w:br")
    break_element.set(qn("w:type"), "page")
    run._r.append(break_element)


def _heading(document: Document, text: str, level: int, style: dict[str, Any]) -> Paragraph:
    paragraph = cast(Paragraph, document.add_heading(text, level=level))
    apply_paragraph_style(paragraph, style)
    for run in paragraph.runs:
        apply_run_style(run, style)
    return paragraph


def heading_1(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    return _heading(document, text, 1, style)


def heading_2(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    return _heading(document, text, 2, style)


def heading_3(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    return _heading(document, text, 3, style)


def body(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    paragraph = document.add_paragraph(text)
    apply_paragraph_style(paragraph, style)
    for run in paragraph.runs:
        apply_run_style(run, style)
    return paragraph


def bullet(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    paragraph = document.add_paragraph()
    apply_paragraph_style(paragraph, style)
    glyph = str(style.get("glyph", "• "))
    run = paragraph.add_run(glyph + text)
    apply_run_style(run, style)
    return paragraph


def bold_lead(document: Document, bold_part: str, rest: str, style: dict[str, Any]) -> Paragraph:
    paragraph = document.add_paragraph()
    apply_paragraph_style(paragraph, style)
    glyph = str(style.get("glyph", "• "))
    run_bold = paragraph.add_run(glyph + bold_part)
    apply_run_style(run_bold, style)
    run_bold.bold = True
    run_rest = paragraph.add_run(rest)
    apply_run_style(run_rest, style)
    run_rest.bold = False
    return paragraph


def reference(document: Document, text: str, style: dict[str, Any]) -> Paragraph:
    paragraph = document.add_paragraph()
    apply_paragraph_style(paragraph, style)
    run = paragraph.add_run(text)
    apply_run_style(run, style)
    return paragraph

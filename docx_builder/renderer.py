from __future__ import annotations

from functools import partial
from typing import Any

from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt

from docx_builder.elements import body, bold_lead, bullet, h1, h2, h3, page_break, reference
from docx_builder.figure import figure as _figure_base
from docx_builder.figure import figure_pair as _figure_pair_base
from docx_builder.styles import StyleResolver
from docx_builder.summary import build_summary

_TEXT_DISPATCHERS = {
    "h1": h1,
    "h2": h2,
    "h3": h3,
    "body": body,
    "bullet": bullet,
    "reference": reference,
}


def fill_cover(document: Document, cover: dict[str, object]) -> None:
    if not document.tables:
        return
    cover_table = document.tables[0]
    rows: list[object] = cover.get("rows", [])  # type: ignore[assignment]
    for index, value in enumerate(rows):
        if index >= len(cover_table.rows):
            break
        cover_table.rows[index].cells[1].paragraphs[0].add_run(str(value))

    ai_declaration = cover.get("ai_declaration")
    if not ai_declaration:
        return
    ai_paragraph = document.add_paragraph()
    ai_paragraph.paragraph_format.space_before = Pt(14)
    bold_run = ai_paragraph.add_run("AI Use Declaration: ")
    bold_run.bold = True
    bold_run.font.size = Pt(11)
    rest_run = ai_paragraph.add_run(str(ai_declaration))
    rest_run.font.size = Pt(11)


def render_sections(
    document: Document,
    sections: list[dict[str, Any]],
    *,
    images_dir: str,
    resolver: StyleResolver,
) -> None:
    figure_call = partial(_figure_base, images_dir=images_dir)
    figure_pair_call = partial(_figure_pair_base, images_dir=images_dir)

    has_any_hidden = any(bool(item.get("hide_page_counter")) for item in sections)
    section_break_added = False

    for item in sections:
        hide = bool(item.get("hide_page_counter", False))

        if has_any_hidden and not hide and not section_break_added:
            document.add_section(WD_SECTION.CONTINUOUS)
            section_break_added = True

        call = str(item["call"])
        inline_style = item.get("style") if isinstance(item.get("style"), dict) else None

        if call == "page_break":
            page_break(document)
            continue

        if call == "toc":
            build_summary(document, levels=str(item.get("levels", "1-2")))
            continue

        if call in _TEXT_DISPATCHERS:
            style = resolver.resolve(call, inline_style)
            _TEXT_DISPATCHERS[call](document, str(item["text"]), style)
            continue

        if call == "bold_lead":
            style = resolver.resolve("bold_lead", inline_style)
            bold_lead(document, str(item["bold"]), str(item["rest"]), style)
            continue

        if call == "figure":
            fig_style = resolver.resolve("figure", inline_style)
            caption_style = resolver.resolve("figure_caption", item.get("caption_style"))
            figure_call(
                document,
                str(item["filename"]),
                str(item["label"]),
                str(item["caption"]),
                figure_style=fig_style,
                caption_style=caption_style,
                width=item.get("width"),
            )
            continue

        if call == "figure_pair":
            fig_style = resolver.resolve("figure_pair", inline_style)
            caption_style = resolver.resolve("figure_caption", item.get("caption_style"))
            figure_pair_call(
                document,
                str(item["filename1"]),
                str(item["filename2"]),
                str(item["label"]),
                str(item["caption"]),
                figure_style=fig_style,
                caption_style=caption_style,
                width1=item.get("width1"),
                width2=item.get("width2"),
            )
            continue

        raise ValueError(f"unknown call: {call!r}")

from __future__ import annotations

from typing import Any

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from lxml.etree import _Element

from docx_builder.styles import apply_run_style, parse_align, parse_length_pt


def _half_point_size(style: dict[str, Any]) -> str:
    size = parse_length_pt(style.get("font_size"))
    if size is None:
        return "18"
    return str(int(round(float(size.pt) * 2)))


def _make_field_run(instruction: str, size_half_points: str) -> _Element:
    run_element = OxmlElement("w:r")

    run_properties = OxmlElement("w:rPr")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), size_half_points)
    run_properties.append(font_size)
    run_element.append(run_properties)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run_element.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run_element.append(instr)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run_element.append(end)

    return run_element


def _make_separator_run(text: str, size_half_points: str) -> _Element:
    run_element = OxmlElement("w:r")

    run_properties = OxmlElement("w:rPr")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), size_half_points)
    run_properties.append(font_size)
    run_element.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)

    return run_element


def _parse_footer_separator(template: str) -> str:
    if "{page}" not in template or "{total}" not in template:
        return " / "
    _prefix, rest = template.split("{page}", 1)
    middle, _suffix = rest.split("{total}", 1)
    return middle


def _add_footer_to_section(section: Section, style: dict[str, Any]) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()

    align = parse_align(style.get("align"))
    if align is not None:
        paragraph.alignment = align

    size_half = _half_point_size(style)
    template = str(style.get("format", "{page} / {total}"))
    middle = _parse_footer_separator(template)

    paragraph_element = paragraph._p
    paragraph_element.append(_make_field_run("PAGE", size_half))
    paragraph_element.append(_make_separator_run(middle, size_half))
    paragraph_element.append(_make_field_run("NUMPAGES", size_half))

    for run in paragraph.runs:
        apply_run_style(run, style)


def _clear_section_footer(section: Section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        footer.paragraphs[0].clear()


def add_page_numbers(
    document: Document,
    *,
    style: dict[str, Any],
    skip_cover_sections: bool = False,
) -> None:
    sections = list(document.sections)
    if not sections:
        return
    if skip_cover_sections:
        _clear_section_footer(sections[0])
        for section in sections[1:]:
            _add_footer_to_section(section, style)
    else:
        for section in sections:
            _add_footer_to_section(section, style)
